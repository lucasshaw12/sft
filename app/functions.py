#! python3
# Functions.py

import requests
import os
import csv
from pathlib import Path
import zipfile
import feedparser
import docx
import re
import send2trash
import pandas as pd
import matplotlib.pyplot as plt
from fpdf import FPDF
import time
import datetime
import sys


##########################################
# Download and organise raw data for UK housing prices
##########################################


def download_uk_housing_data(raw_csv_data: str) -> str:
    """Download then save UK housing price data.
    :param raw_csv_data: the url which contains the .csv file of raw data.
    :type raw_csv_data: str
    :returns: str of the downloaded filename
    """

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    res = requests.get(raw_csv_data)
    res.raise_for_status()

    # create folder to store the data
    os.makedirs('../raw data/housingprices', exist_ok=True)

    # Download the .csv file data
    raw_data_file = open(os.path.join('../raw data/housingprices', 'uk_average_house_price.csv'), 'wb')
    for chunk in res.iter_content(100000):
        raw_data_file.write(chunk)
    raw_data_file.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'downloaded UK housing price data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()
    return raw_csv_data


def organise_uk_housing_data():
    """Parse, clean then organise csv data.
    Specific to this data only, as raw .csv is unique."""

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    # Read the raw csv  data
    with open('../raw data/housingprices/uk_average_house_price.csv') as csv_file:
        dict_reader = csv.DictReader(csv_file, ['Date', 'England (pop)', 'Wales (pop)', 'Scotland (pop)',
                                                'NI (pop)'])  # Create own headers
        csv_rows = []
        for row in dict_reader:
            if dict_reader.line_num < 11:  # remove irrelevant rows of data from file
                continue
            csv_rows.append(row)

    # Write out the csv file to a clean file
    csv_obj = open(os.path.join('../raw data/housingprices', 'clean_uk_average_house_price.csv'), 'w', newline='')
    csv_writer = csv.DictWriter(csv_obj, ['Date', 'England (pop)', 'Wales (pop)', 'Scotland (pop)', 'NI (pop)'])
    csv_writer.writeheader()
    for row in csv_rows:
        csv_writer.writerow(row)
    csv_obj.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'organised UK housing price data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()


##########################################
# END
##########################################

##########################################
# Download and organise raw data for UK electricity cpi index
##########################################


def download_uk_electric_cpi_change_data(raw_csv_data: str) -> str:
    """Download UK electricity CPI (Consumer Price Index 2015=100) price data
    :param raw_csv_data: the url which contains the .csv file of raw data.
    :type raw_csv_data: str
    :returns: str of the downloaded filename
    """

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    res = requests.get(raw_csv_data)
    res.raise_for_status()  # Use to time how long the function takes to complete

    # create folder to store the data
    os.makedirs('../raw data/electricdata', exist_ok=True)

    # Download the .csv file data
    raw_data_file = open(os.path.join('../raw data/electricdata', 'uk_cpi_electric_index.csv'), 'wb')
    for chunk in res.iter_content(100000):
        raw_data_file.write(chunk)
    raw_data_file.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'downloaded UK electricity CPI price data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()
    return raw_csv_data


def organise_uk_electric_cpi_change_data():
    """Parse, clean then organise csv data.
    Specific to this data only, as raw .csv is unique."""

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    # Read the raw csv  data
    with open('../raw data/electricdata/uk_cpi_electric_index.csv') as csv_file:
        dict_reader = csv.DictReader(csv_file, ['Date', 'Value'])  # Create own headers
        csv_rows = []
        for row in dict_reader:
            if dict_reader.line_num < 9 or dict_reader.line_num > 42:  # remove irrelevant rows of data from file
                continue
            csv_rows.append(row)

    # Write out the csv file to a clean file
    csv_obj = open(os.path.join('../raw data/electricdata', 'clean_uk_cpi_electric_index.csv'), 'w', newline='')
    csv_writer = csv.DictWriter(csv_obj, ['Date', 'Value'])
    csv_writer.writeheader()
    for row in csv_rows:
        csv_writer.writerow(row)
    csv_obj.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'organised UK electricity CPI price data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()


##########################################
# END
##########################################

##########################################
# Download and organise raw data for UK gas cpi index
##########################################

def download_uk_gas_cpi_change_price(raw_csv_data):
    """Download UK gas CPI (Consumer Price Index 2015=100) price data
    :param raw_csv_data: the url which contains the .csv file of raw data.
    :type raw_csv_data: str
    :returns: str of the downloaded filename
    """

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    res = requests.get(raw_csv_data)
    res.raise_for_status()  # Use to time how long the function takes to complete

    # create folder to store the data
    os.makedirs('../raw data/gasdata', exist_ok=True)

    # Download the .csv file data
    raw_data_file = open(os.path.join('../raw data/gasdata', 'uk_cpi_gas_index.csv'), 'wb')
    for chunk in res.iter_content(100000):
        raw_data_file.write(chunk)
    raw_data_file.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'downloaded UK gas CPI price data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()
    return raw_csv_data


def organise_uk_gas_cpi_change_data():
    """Parse, clean then organise csv data.
    Specific to this data only, as raw .csv is unique."""

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    # Read the raw csv  data
    with open('../raw data/gasdata/uk_cpi_gas_index.csv') as csv_file:
        dict_reader = csv.DictReader(csv_file, ['Date', 'Value'])  # Create own headers
        csv_rows = []
        for row in dict_reader:
            if dict_reader.line_num < 9 or dict_reader.line_num > 42:  # remove irrelevant rows of data from file
                continue
            csv_rows.append(row)

    # Write out the csv file to a clean file
    csv_obj = open(os.path.join('../raw data/gasdata', 'clean_uk_cpi_gas_index.csv'), 'w', newline='')
    csv_writer = csv.DictWriter(csv_obj, ['Date', 'Value'])
    csv_writer.writeheader()
    for row in csv_rows:
        csv_writer.writerow(row)
    csv_obj.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'organised UK gas CPI price data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()


##########################################
# END
##########################################

##########################################
# Download and organise raw data for UK water supply cpi index
##########################################


def download_uk_water_cpi_change_data(raw_csv_data):
    """Download then save UK water cpi data
    :param raw_csv_data: the url which contains the .csv file of raw data.
    :type raw_csv_data: str
    :returns: str of the downloaded filename
    """

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    res = requests.get(raw_csv_data)
    res.raise_for_status()

    # create folder to store the data
    os.makedirs('../raw data/watersupplydata', exist_ok=True)

    # Download the .csv file data
    raw_data_file = open(os.path.join('../raw data/watersupplydata', 'uk_cpi_water_index.csv'), 'wb')
    for chunk in res.iter_content(100000):
        raw_data_file.write(chunk)
    raw_data_file.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'downloaded UK water supply data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()
    return raw_csv_data


def organise_uk_water_cpi_change_data():
    """Parse, clean then organise csv data.
    Specific to this data only, as raw .csv is unique."""

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    # Read the raw csv  data
    with open('../raw data/watersupplydata/uk_cpi_water_index.csv') as csv_file:
        dict_reader = csv.DictReader(csv_file, ['Date', 'Value'])  # Create own headers
        csv_rows = []
        for row in dict_reader:
            if dict_reader.line_num < 9 or dict_reader.line_num > 42:  # remove irrelevant rows of data from file
                continue
            csv_rows.append(row)

    # Write out the csv file to a clean file
    csv_obj = open(os.path.join('../raw data/watersupplydata', 'clean_uk_cpi_water_index.csv'), 'w', newline='')
    csv_writer = csv.DictWriter(csv_obj, ['Date', 'Value'])
    csv_writer.writeheader()
    for row in csv_rows:
        csv_writer.writerow(row)
    csv_obj.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'organised UK water supply data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()


##########################################
# END
##########################################

##########################################
# Download and organise raw data for UK consumer price index CPI
##########################################


def download_uk_cpi_data(raw_csv_data):
    """Download then save UK water cpi data
    :param raw_csv_data: the url which contains the .csv file of raw data.
    :type raw_csv_data: str
    :returns: str of the downloaded filename
    """
    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    res = requests.get(raw_csv_data)
    res.raise_for_status()

    # create folder to store the data
    os.makedirs('../raw data/cpi', exist_ok=True)

    # Download the .zip file data
    raw_data_file = open(os.path.join('../raw data/cpi', 'uk_cpi.zip'), 'wb')
    for chunk in res.iter_content(100000):
        raw_data_file.write(chunk)
    raw_data_file.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'downloaded UK Consumer price index (CPI) data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()
    return raw_csv_data


def organise_uk_cpi_data():
    """Unzip, parse, clean then organise csv data.
    Specific to this data only, as raw .csv is unique."""

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete
    p = Path('../raw data/cpi/')

    # Unzip the .zip file
    data_zip = zipfile.ZipFile(p / 'uk_cpi.zip')
    data_zip.extractall(p)

    # Organise files from the downloaded then extracted 'ukcpi.zip' file
    for folder, subfolder, files in os.walk(p):
        for file in files:
            if file.startswith('API'):
                uk_cpi_file = os.path.join(folder, file)

            if file.startswith('Meta'):
                filepath = os.path.join(folder, file)
                print(f'Deleted unused files {filepath}')
                send2trash.send2trash(filepath)  # Send unused files to bin

    # Read the csv file
    csv_rows = []
    with open(uk_cpi_file) as csv_file:
        reader_obj = csv.reader(csv_file)
        for row in reader_obj:
            if reader_obj.line_num != 87 and reader_obj.line_num != 5:
                continue  # Skip all other rows besides 5 and 87
            csv_rows.append(row)
        csv_file.close()

        # Write out the csv file
        csv_file_obj = open(os.path.join(p, 'uk_cpi.csv'), 'w', newline='')
        csv_writer = csv.writer(csv_file_obj)
        for row in csv_rows:
            csv_writer.writerow(row)
        csv_file_obj.close()

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'unzipped and organised UK CPI data - DONE - time taken = {end_time}'.upper())
    sys.stdout.flush()


##########################################
# END
##########################################

##########################################
# Download and organise T212 personal holdings
##########################################

# def download_t212_data():
#     pass
#
#
# def organise_t212_data():
#     pass

##########################################
# END
##########################################

##################################################
# Manipulating .csv files to .xlsx files
##################################################

def convert_csv_to_excel(csv_file: str):
    """Convert .csv files to .xlsx in preparation for chart plotting.
    :param csv_file: The file location of the spreadsheet (.csv) file.
    """

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    df = pd.read_csv(csv_file)
    xlsx_file = csv_file.rstrip('.csv') + '.xlsx'  # create filename for xlsx file
    df.to_excel(xlsx_file, index=None, header=True)  # save xlsx file

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'converted {os.path.basename(csv_file)} to .xlsx format - time taken = {end_time}'.upper())
    sys.stdout.flush()


def transpose_xlsx(xlsx_file_to_transpose: str):
    """Used on the 'ukcpi.csv' file due to horizontal format - created for matplotlib chart
    Transpose, create a header then remove unused rows.
    :param xlsx_file_to_transpose: The file location of the spreadsheet (.xlsx) file."""

    df = pd.read_excel(xlsx_file_to_transpose, header=None)
    df = df.T  # Transpose the dataframe
    df.columns = df.iloc[0]  # Cut the first row
    df.columns = ['Date', 'Value']  # Create column headers
    df = df[4:]  # ignore the first few rows
    df.to_excel('../raw data/cpi/clean_ukcpi.xlsx', index=False)


##########################################
# END
##########################################

##########################################
# Plotting .xlsx files to chart/graphs
##########################################

def plot_graph_from_excel(xlsx_file_to_plot: str):
    """Plot a graph using a .xlsx file.
    :param xlsx_file_to_plot: The file location of the spreadsheet (.xlsx) which shall be plotted."""

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete
    df = pd.read_excel(xlsx_file_to_plot)
    df.head()

    os.makedirs('../charts/', exist_ok=True)

    # Chart header name
    filename = os.path.basename(xlsx_file_to_plot).rstrip('xlsx').lstrip('clean_')
    file_string = re.sub(r'(?<=[a-z])_(?=[a-z])', ' ', filename)

    plt.rcParams['axes.labelsize'] = 8
    plt.rcParams['axes.titlesize'] = 8

    plt.figure(figsize=(20, 9), dpi=100)  # Set chart size
    df.plot.line('Date')
    plt.title(file_string.title(), fontsize=10)
    plt.ylabel(f'{file_string.title()} (£)')  # y axis label
    plt.annotate(f'Sourced from ons.gov.uk. Last updated on {time.ctime()}', (0, 0), (-80, -80),
                 fontsize=8,
                 xycoords='axes fraction', textcoords='offset points', va='top')  # to display data source
    plt.grid(True, color='k', linestyle=':')
    plt.xticks(rotation=90, fontsize=8)

    chart_filename = filename.rstrip('xlsx') + 'png'  # create filename for chart image
    plt.savefig(os.path.join('../charts/', os.path.basename(chart_filename)), bbox_inches='tight')

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(
        f'plotted graph from {os.path.basename(xlsx_file_to_plot)} to {os.path.basename(chart_filename)}- time '
        f'taken = {end_time}'.upper())
    sys.stdout.flush()


def display_stock_chart_pdf():
    """Add the stock chart .png files to a single PDF for ease of viewing."""

    start_time = datetime.datetime.now()
    pdf = FPDF()

    # Find all .png files
    image_list = []
    for folder, subfolder, files in os.walk('../charts'):
        for file in files:
            if file.endswith('.png'):
                png_filename = os.path.join(folder, file)
                image_list.append(png_filename)  # append all images to image_list

    for image in image_list:
        pdf.add_page()
        pdf.image(image, 10, 10, 150, 150)
    pdf.output("../charts/uk_economic_data.pdf", "F")

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(f'pdf charts complete from {os.path.basename(str(files[1:]))}- time taken = {end_time}'.upper())
    sys.stdout.flush()


##########################################
# END
##########################################

##########################################
# RSS Feeds
##########################################


def download_bbc_rss_feed(rss_feed_url: str):
    """Download RSS feed data from the BBC and write to a .docx file.
    Used for all BBC RSS feed categories ie education, politics, health etc.
    :param rss_feed_url: url of the RSS feed location
    """

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete
    last_updated_time = time.ctime()

    data = feedparser.parse(rss_feed_url)
    i = 0

    # Create word doc
    doc = docx.Document()

    # Create news folder
    os.makedirs('../newsfeeds/bbcnews', exist_ok=True)

    # Regular expression to provide respective topic name to word doc header and filename
    topic_regex = re.compile(r'''(
    	((http|https)://[A-Za-z.]+/[A-Za-z]+/) # pre topic url 'http://feeds.bbci.co.uk/news/'
    	([A-Za-z_]+) # topic name in url i.e 'business'
    	(/\w{3}\.\w{3})  # post topic url '/rss.xml'
    	)''', re.VERBOSE)
    topic_mo = topic_regex.search(rss_feed_url)
    topic_name = topic_mo.group(4)  # Topic name to insert into filename

    doc.add_heading('BBC ' + topic_name.title() + ' News', 1)
    doc.add_paragraph(f'Last updated on: {last_updated_time}')
    while i < 10:  # Find first 10 articles
        # Store text variable to write to word document
        feed_title_str = data['entries'][i]["title"]
        feed_description_str = data['entries'][i]['description']
        feed_description_link = data['entries'][i]['link']

        doc.add_paragraph(feed_title_str, 'Heading 3')
        doc.add_paragraph(feed_description_str)
        doc.add_paragraph(feed_description_link, 'Normal')
        doc.add_paragraph()

        i = i + 1  # Move to the next article

    doc.save(f'../newsfeeds/bbcnews/bbc_news_{topic_name}.docx')

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(
        f'current bbc news articles saved to ../newsfeeds/bbcnews/bbc_news_{topic_name}.docx - time taken = {end_time}'.upper())
    sys.stdout.flush()


def download_investing_rss_feed(rss_feed_url):
    """Download RSS feed data from investing.com and write to a .docx file.
    :param rss_feed_url: url of the RSS feed location."""

    start_time = datetime.datetime.now()  # Use to time how long the function takes to complete

    data = feedparser.parse(rss_feed_url)
    i = 0

    # Create word doc
    doc = docx.Document()

    # Create news folder
    os.makedirs('../newsfeeds/investing.com', exist_ok=True)

    while i < 10:  # Find first 10 articles
        # Store text variable to write to word document
        feed_title_str = data['entries'][i]["title"]
        feed_pubdate = data['entries'][i]['published']
        feed_link = data['entries'][i]['link']

        doc.add_paragraph(feed_title_str, 'Heading 3')
        doc.add_paragraph('Published on ' + feed_pubdate, 'Normal')
        doc.add_paragraph(feed_link, 'Normal')
        doc.add_paragraph()

        i = i + 1  # Move to the next article

    doc.save(f'../newsfeeds/investing.com/investingnews.docx')

    end_time = datetime.datetime.now() - start_time
    sys.stdout.write(
        f'current investing.com news articles saved to ../newsfeeds/investing.com/investingnews.docx - time taken = {end_time}'.upper())
    sys.stdout.flush()

# def download_morningstar_rss_feed(rss_feed_url):
#     # Download RSS feed data from Morningstar and write to a .docx file
#     start_time = datetime.datetime.now()  # Use to time how long the function takes to complete
#
#     data = feedparser.parse(rss_feed_url)
#     i = 0
#
#     # Create word doc
#     doc = docx.Document()
#
#     # Create news folder
#     os.makedirs('../newsfeeds/morningstar', exist_ok=True)
#
#     while i < 10:  # Find first 10 articles
#         # Store text variable to write to word document
#         feed_title = data['entries'][i]["title"]
#         # feed_description = data['entries'][i]['description']
#         feed_link = data['entries'][i]['link']
#         feed_pubdate = data['entries'][i]['published']
#
#         doc.add_paragraph(feed_title, 'Heading 3')
#         # doc.add_paragraph(feed_description, 'Normal')
#         doc.add_paragraph(feed_pubdate, 'Normal')
#         doc.add_paragraph(feed_link, 'Normal')
#         doc.add_paragraph()
#
#         i = i + 1  # Move to the next article
#
#     doc.save(f'../newsfeeds/morningstar/investingnews.docx')
#
#     end_time = datetime.datetime.now() - start_time
#     print(f'current morningstar news articles saved to ../newsfeeds/morningstar/morningstar.docx - time taken = {end_time}'.upper())


##########################################
# END
##########################################
